import json
import logging
import re
from typing import List, Dict, Any, Optional

from docx import Document
from docx.shared import Inches
from pydantic import BaseModel, Field, ValidationError

from ..core.config import settings
from .llm_service import get_llm_response
from .vector_store_service import vector_store

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# 1. Pydantic Models for structured data
class OutlineSection(BaseModel):
    section_title: str = Field(..., description="The title of this section.")
    description: str = Field(..., description="A brief description of what this section will cover.")

class ReviewOutline(BaseModel):
    review_title: str = Field(..., description="The overall title of the literature review.")
    introduction: str = Field(..., description="A brief paragraph describing the introduction's scope.")
    body: List[OutlineSection] = Field(..., description="The main sections of the review body.")
    conclusion: str = Field(..., description="A brief paragraph describing the conclusion's scope.")

class LiteratureReviewService:
    """
    A service to generate a literature review in a multi-step, controllable process.
    """

    def _gather_context(self, topic: str, max_results: int = 50) -> Dict[str, Any]:
        """Gathers context from the vector store based on a topic."""
        logging.info(f"Expanding topic '{topic}' and gathering context...")
        expansion_prompt = f'''
        Given the research topic \"{topic}\", generate 3 to 5 diverse and specific questions that a researcher might ask.
        These questions should cover different facets of the topic. Return the questions as a numbered list.
        '''
        sub_queries_str = get_llm_response(expansion_prompt, use_reasoner=True)
        sub_queries = [q.split('.', 1)[1].strip() for q in sub_queries_str.split('\n') if q.strip() and q.strip()[0].isdigit()]
        queries_to_run = [topic] + (sub_queries or [])
        
        all_context_data = {}
        n_per_query = max(1, max_results // len(queries_to_run))

        for query in queries_to_run:
            results = vector_store.query(
                query_text=query,
                n_results=n_per_query,
                collection_name=settings.SUMMARIES_COLLECTION_NAME
            )
            if results and results.get('documents'):
                docs = results['documents'][0]
                metas = results['metadatas'][0]
                for doc, meta in zip(docs, metas):
                    if doc not in all_context_data:
                        all_context_data[doc] = meta
        
        logging.info(f"Aggregated {len(all_context_data)} unique document chunks.")
        return all_context_data

    def generate_outline(self, topic: str) -> Dict[str, Any]:
        """
        Step 1: Generates a structured outline for the literature review.
        Returns the outline and the context documents needed for the next step.
        """
        context_docs = self._gather_context(topic)
        if not context_docs:
            return {"error": "Could not find relevant documents to generate an outline."}

        context_summary = "\n".join([f"- {doc[:200]}..." for doc in list(context_docs.keys())[:10]])

        prompt = f'''
        You are a research strategist. Based on the following document summaries, create a structured and logical outline for a literature review on the topic: "{topic}".

        DOCUMENT SUMMARIES:
        {context_summary}

        INSTRUCTIONS:
        Respond with a single, valid JSON object that strictly follows this Pydantic model:
        ```json
        {{
            "review_title": "<A concise and informative title>",
            "introduction": "<A one-sentence description of the introduction's goal>",
            "body": [
                {{
                    "section_title": "<Title of the first thematic section>",
                    "description": "<A one-sentence summary of this section's content>"
                }}
            ],
            "conclusion": "<A one-sentence description of the conclusion's goal>"
        }}
        ```
        '''
        try:
            outline_str = get_llm_response(prompt, json_mode=True, use_reasoner=True)
            outline = ReviewOutline.model_validate_json(outline_str)
            return {
                "outline": outline.model_dump(),
                "context_docs": context_docs
            }
        except (ValidationError, json.JSONDecodeError) as e:
            logging.error(f"Failed to generate or validate review outline: {e}")
            return {"error": f"Failed to generate a valid outline from LLM. Details: {e}"}

    def _format_apa_citation(self, metadata: Dict[str, Any], source_file: str) -> str:
        """Formats a citation in a simplified APA-like style."""
        authors = metadata.get('authors', 'N/A').split('; ')[0] + " et al." if ';' in metadata.get('authors', '') else metadata.get('authors', 'N/A')
        year = metadata.get('publication_year', 'n.d.')
        title = metadata.get('title', source_file)
        journal = metadata.get('journal', 'Source not specified')
        return f"{authors} ({year}). {title}. *{journal}*."

    def generate_review_from_outline(self, outline: Dict[str, Any], context_docs: Dict[str, Any]) -> Dict[str, Any]:
        """
        Step 2: Generates the full review content based on a user-approved outline.
        """
        try:
            validated_outline = ReviewOutline.model_validate(outline)
        except ValidationError as e:
            return {"error": f"Invalid outline structure provided. Details: {e}"}

        # Create citation map: Source_1 -> metadata
        source_map: Dict[str, Dict[str, Any]] = {}
        # Create reverse map: doc_content -> citation_tag
        doc_to_citation: Dict[str, str] = {}
        for i, (doc, meta) in enumerate(context_docs.items(), 1):
            tag = f"Source_{i}"
            source_map[tag] = meta
            doc_to_citation[doc] = tag

        context_str = "\n\n---\n\n".join([f"Context from [{doc_to_citation[doc]}]:\n{doc}" for doc in context_docs])

        # Generate content for each section
        generated_content = {}
        # Create OutlineSection objects for Introduction and Conclusion
        intro_section = OutlineSection(section_title="Introduction", description="Introduction section")
        conclusion_section = OutlineSection(section_title="Conclusion", description="Conclusion section")
        
        for section in [intro_section] + validated_outline.body + [conclusion_section]:
            section_title = section.section_title
            logging.info(f"Generating content for section: {section_title}")
            
            section_prompt = f'''
            You are a research writer. Your task is to write the "{section_title}" section of a literature review on "{validated_outline.review_title}".
            Base your writing *ONLY* on the provided CONTEXT. Cite sources using their tags (e.g., [Source_1]) at the end of each sentence where information is used.
            Focus only on writing the content for this specific section. Do not write any other sections.

            CONTEXT:
            ---
            {context_str}
            ---

            Now, write the "{section_title}" section:
            '''
            section_content = get_llm_response(section_prompt, use_reasoner=True)
            generated_content[section_title] = section_content.strip()

        # Dynamically build reference list
        used_tags = set(re.findall(r'\[(Source_\d+)\]', " ".join(generated_content.values())))
        references = {tag: self._format_apa_citation(source_map[tag], tag) for tag in sorted(used_tags) if tag in source_map}

        return {
            "title": validated_outline.review_title,
            "content": generated_content,
            "references": references
        }

    def export_review_to_docx(self, review_data: Dict[str, Any]) -> Any:
        """
        Step 3: Exports the generated review data to a .docx file.
        """
        doc = Document()
        doc.add_heading(review_data.get('title', 'Literature Review'), level=0)

        content = review_data.get('content', {})
        
        # Add Introduction
        if 'Introduction' in content:
            doc.add_heading('Introduction', level=1)
            doc.add_paragraph(content['Introduction'])

        # Add Body Sections
        for section_title, section_text in content.items():
            if section_title not in ['Introduction', 'Conclusion']:
                doc.add_heading(section_title, level=1)
                doc.add_paragraph(section_text)

        # Add Conclusion
        if 'Conclusion' in content:
            doc.add_heading('Conclusion', level=1)
            doc.add_paragraph(content['Conclusion'])

        # Add References
        references = review_data.get('references', {})
        final_ref_map = {}  # Initialize empty dict
        if references:
            doc.add_heading('References', level=1)
            # Renumber references for final output [1], [2], ...
            final_ref_map = {tag: f"[{i}]" for i, tag in enumerate(references.keys(), 1)}
            for tag, formatted_ref in references.items():
                p = doc.add_paragraph(style='List Number')
                p.add_run(formatted_ref)
                # Replace in-text citations
                for section_title, section_text in content.items():
                    content[section_title] = section_text.replace(f'[{tag}]', f' {final_ref_map[tag]}')

        # Re-write paragraphs with updated citations (this is inefficient but simple)
        for p in doc.paragraphs:
            for tag, num_ref in final_ref_map.items():
                if f'[{tag}]' in p.text:
                    # This is a simplification; robust replacement is complex.
                    # A better approach would be to rebuild the document.
                    pass # Placeholder for more complex replacement logic

        logging.info("Successfully created .docx document object.")
        return doc

# Singleton instance
literature_review_service = LiteratureReviewService()
