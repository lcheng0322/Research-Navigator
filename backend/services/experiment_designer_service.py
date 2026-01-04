import json
import logging
from typing import Any, Dict
from docx import Document

from ..schemas.experiment_schemas import (
    ExperimentDesignRequest, 
    ExperimentDesign, 
    Hypothesis,
    DesignSession,
    CreateSessionResponse,
    Critique
)
from ..services.vector_store_service import vector_store
from ..services.llm_service import get_llm_response
from ..core.config import settings

logger = logging.getLogger(__name__)

# In-memory storage for design sessions. 
# For production, this would be replaced with a database or a Redis cache.
design_sessions: Dict[str, DesignSession] = {}

class ExperimentDesignerService:
    """
    Service for the interactive Experimental Design Planner.
    """

    def _format_context(self, context_docs: list[dict[str, Any]]) -> str:
        """Formats the retrieved documents into a string for the LLM prompt."""
        if not context_docs:
            return "No relevant information found in the knowledge base."
        
        formatted_parts = ["Relevant information from the knowledge base:"]
        for i, doc in enumerate(context_docs):
            metadata = doc.get('metadata', {})
            source = metadata.get('source', 'Unknown')
            page = metadata.get('page_number', 'N/A')
            # Prefer 'text' if present; otherwise fall back to 'content'
            content = (doc.get('text') or doc.get('content') or '')
            
            formatted_parts.append(
                f"\n--- Context {i+1} (Source: {source}, Page: {page}) ---\n{content}"
            )
        
        return "".join(formatted_parts)

    def _retrieve_context(self, topic: str) -> tuple[str, int]:
        """Retrieves and formats context from the vector store."""
        try:
            logger.info(f"Retrieving context from vector store for topic: '{topic}'...")
            retrieved_results = vector_store.query(
                query_text=topic,
                n_results=5,
                collection_name=settings.CHUNKS_COLLECTION_NAME
            )
            
            context_docs = []
            if retrieved_results and retrieved_results.get('documents') and retrieved_results.get('metadatas'):
                docs = retrieved_results['documents'][0]
                metas = retrieved_results['metadatas'][0]
                if docs and metas:
                    for doc, meta in zip(docs, metas):
                        context_docs.append({"text": doc, "metadata": meta})

            formatted_context = self._format_context(context_docs)
            logger.info(f"Retrieved {len(context_docs)} documents for context.")
            return formatted_context, len(context_docs)
        except Exception as e:
            logger.error(f"Failed to retrieve context from vector store: {e}", exc_info=True)
            return "Warning: Could not retrieve information from the knowledge base due to an error.", 0

    def create_session(self, request: ExperimentDesignRequest) -> CreateSessionResponse:
        """
        Creates a new design session, retrieves context, and generates a preliminary hypothesis.
        """
        logger.info(f"Creating new design session for topic: '{request.research_topic}'...")

        # 1. Retrieve context
        formatted_context, doc_count = self._retrieve_context(request.research_topic)

        # 2. Generate preliminary hypothesis
        hypothesis_prompt = f"""\
        Based on the following research topic and context from a knowledge base, formulate a clear, testable research hypothesis and a brief summary of the context that informed it.

        **Research Topic:** {request.research_topic}

        **Relevant Context:**
        {formatted_context}

        **Your Task:**
        Return a single, valid JSON object with two keys:
        1. `hypothesis_text`: The specific, testable hypothesis.
        2. `context_summary`: A one-sentence summary of the key information from the context that justifies this hypothesis.
        """
        
        hypothesis_response_str = get_llm_response(hypothesis_prompt, json_mode=True, use_reasoner=False)
        hypothesis_json = json.loads(hypothesis_response_str)
        hypothesis = Hypothesis.model_validate(hypothesis_json)

        # 3. Create and store the session
        session = DesignSession(
            request=request,
            retrieved_context=formatted_context,
            hypothesis=hypothesis
        )
        design_sessions[session.session_id] = session
        
        logger.info(f"Successfully created session {session.session_id} with initial hypothesis.")

        return CreateSessionResponse(session_id=session.session_id, hypothesis=session.hypothesis)

    def generate_full_design(self, session_id: str) -> ExperimentDesign:
        """
        Generates the full, detailed experimental design for an existing session.
        """
        if session_id not in design_sessions:
            raise ValueError(f"Design session with ID '{session_id}' not found.")

        session = design_sessions[session_id]
        logger.info(f"Generating full design for session {session_id}...")

        # Generate the final, detailed design using a structured prompt
        schema_json_string = json.dumps(ExperimentDesign.model_json_schema(), indent=2)

        final_design_prompt = f"""\
        You are an expert research assistant. Your task is to expand a given hypothesis into a full experimental design, based on the original user request and the context provided.

        **Original User Request:**
        - Research Topic: {session.request.research_topic}
        - Key Variables: {session.request.variables}
        - Constraints: {session.request.constraints}

        **Approved Hypothesis:** {session.hypothesis.hypothesis_text}

        **Relevant Context from Knowledge Base:**
        {session.retrieved_context}

        **Your Task:**
        Generate a comprehensive and scientifically rigorous experimental design. You MUST output a single, valid JSON object that strictly adheres to the following schema. Do not include any text or explanations outside of the JSON object.

        **JSON Schema:**
        {schema_json_string}
        """
        
        llm_response_str = get_llm_response(final_design_prompt, json_mode=True, use_reasoner=False)

        # Parse and validate the response
        try:
            llm_json = json.loads(llm_response_str)
            design = ExperimentDesign.model_validate(llm_json)
            
            # Update and store the final design in the session object
            session.final_design = design
            session.status = "design_completed"
            design_sessions[session.session_id] = session
            
            logger.info(f"Successfully generated and validated final design for session {session_id}.")
            return design
        except (json.JSONDecodeError, TypeError) as e:
            logger.error(f"Failed to decode LLM JSON response for session {session_id}: {e}\nResponse was: {llm_response_str}", exc_info=True)
            raise ValueError(f"The LLM returned an invalid JSON object. Raw response: {llm_response_str}")
        except Exception as e:
            logger.error(f"Failed to validate LLM response for session {session_id}: {e}", exc_info=True)
            raise ValueError(f"The LLM response did not match the required data structure. Raw response: {llm_response_str}")

    def review_and_refine_design(self, session_id: str) -> ExperimentDesign:
        """
        Adds a review step to the design process, where a critic agent suggests
        improvements, and a refiner agent incorporates them.
        """
        if session_id not in design_sessions or not design_sessions[session_id].final_design:
            raise ValueError(f"A completed design for session ID '{session_id}' must be generated before it can be refined.")

        session = design_sessions[session_id]
        original_design = session.final_design
        
        # Additional type check to ensure original_design is not None
        if original_design is None:
            raise ValueError(f"No final design found for session ID '{session_id}'")
            
        logger.info(f"Starting review and refinement for design in session {session_id}...")

        # 1. Get critique from a "critic" agent
        critic_prompt = f"""\
        You are a skeptical, rigorous, and experienced scientific reviewer. Your task is to find potential flaws and suggest improvements for the following experimental design.

        **Experimental Design to Review:**
        ```json
        {original_design.model_dump_json(indent=2)}
        ```

        **Your Task:**
        Provide a critical review of this design. Focus on identifying potential flaws, biases, missing controls, or areas where the methodology could be more robust. Output a single, valid JSON object with two keys:
        1. `potential_flaws`: A list of strings, where each string is a specific, identified flaw.
        2. `suggested_improvements`: A list of strings, where each string is a concrete, actionable suggestion for improvement.
        """
        
        critique_response_str = get_llm_response(critic_prompt, json_mode=True, use_reasoner=False)
        critique = Critique.model_validate_json(critique_response_str)
        logger.info(f"Generated critique for session {session_id}: {critique.model_dump_json(indent=2)}")

        # 2. Refine the design based on the critique
        refine_prompt = f"""\
        You are an expert research assistant. Your task is to revise an experimental design based on a critical review.

        **Original Experimental Design:**
        ```json
        {original_design.model_dump_json(indent=2)}
        ```

        **Critical Review and Suggestions:**
        - **Potential Flaws Found:** {critique.potential_flaws}
        - **Suggested Improvements:** {critique.suggested_improvements}

        **Your Task:**
        Generate a new, improved version of the experimental design that addresses the points raised in the review. The revised design should be more robust, clear, and scientifically sound. You MUST output a single, valid JSON object that strictly adheres to the original schema. Do not include any text or explanations outside of the JSON object.
        """

        refined_design_str = get_llm_response(refine_prompt, json_mode=True, use_reasoner=False)
        refined_design = ExperimentDesign.model_validate_json(refined_design_str)

        # 3. Update the session with the refined design
        session.final_design = refined_design
        design_sessions[session_id] = session
        logger.info(f"Successfully refined the design for session {session_id}.")

        return refined_design

    def get_final_design(self, session_id: str) -> ExperimentDesign:
        """
        Retrieve the final (possibly refined) experimental design for a session.
        """
        if session_id not in design_sessions:
            raise ValueError(f"Design session with ID '{session_id}' not found.")
        session = design_sessions[session_id]
        if not session.final_design:
            raise ValueError(f"Final design for session ID '{session_id}' is not available. Please generate the design first.")
        return session.final_design

    def export_design_to_docx(self, design: ExperimentDesign) -> Any:
        """
        Export an ExperimentDesign to a .docx Document object.
        """
        doc = Document()

        # Title
        doc.add_heading(design.title or "Experimental Design", level=0)

        # Core sections
        doc.add_heading("Hypothesis", level=1)
        doc.add_paragraph(design.hypothesis)

        doc.add_heading("Methodology", level=1)
        doc.add_paragraph(design.methodology)

        doc.add_heading("Materials", level=1)
        if design.materials:
            for item in design.materials:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(item)
        else:
            doc.add_paragraph("N/A")

        doc.add_heading("Groups", level=1)
        doc.add_paragraph(f"Control Group: {design.control_group}")
        doc.add_paragraph(f"Experimental Group: {design.experimental_group}")

        doc.add_heading("Protocol Steps", level=1)
        for step in design.steps:
            doc.add_heading(f"Step {step.step_number}", level=2)
            doc.add_paragraph(step.description)
            if step.materials_needed:
                doc.add_paragraph("Materials Needed:")
                for m in step.materials_needed:
                    sp = doc.add_paragraph(style='List Bullet')
                    sp.add_run(m)

        doc.add_heading("Data Analysis Plan", level=1)
        doc.add_paragraph(design.data_analysis_plan)

        if design.potential_risks:
            doc.add_heading("Potential Risks", level=1)
            doc.add_paragraph(design.potential_risks)

        return doc

# Create a singleton instance of the service
experiment_designer_service = ExperimentDesignerService()