from typing import Any, Dict, List
from docx import Document
from docx.shared import Pt


class QueryExportService:
    def export_query_result_to_docx(self, data: Dict[str, Any], query_text: str) -> Document:
        doc = Document()

        # Title
        doc.add_heading('Smart Q&A Result', level=1)
        doc.add_paragraph(f'Query: {query_text}')

        # Synthesized Answer and Confidence
        ra = (data or {}).get('reasoned_answer') or {}
        result = ra.get('result') or {}
        synthesized_answer = result.get('synthesized_answer') or 'N/A'
        limitations = result.get('limitations_analysis') or 'N/A'
        confidence = result.get('confidence_score')
        alternatives: List[str] = result.get('alternative_hypotheses') or []

        doc.add_heading('Synthesized Answer', level=2)
        doc.add_paragraph(synthesized_answer)

        doc.add_heading('Confidence Score', level=2)
        doc.add_paragraph(f'{confidence if confidence is not None else "N/A"}')

        doc.add_heading('Limitations', level=2)
        doc.add_paragraph(limitations)

        if alternatives:
            doc.add_heading('Alternative Hypotheses', level=2)
            for alt in alternatives:
                doc.add_paragraph(alt, style='List Bullet')

        # Query Analysis
        qa = (data or {}).get('query_analysis') or {}
        doc.add_heading('Query Analysis', level=2)
        doc.add_paragraph(f"Intent: {qa.get('intent', 'N/A')}")
        doc.add_paragraph(f"Complexity: {qa.get('complexity', 'N/A')}")
        doc.add_paragraph(f"Domain: {qa.get('domain', 'N/A')}")
        doc.add_paragraph(f"Rewritten Query: {qa.get('rewritten_query', 'N/A')}")
        entities = qa.get('entities') or []
        if entities:
            doc.add_paragraph('Entities:')
            for ent in entities:
                doc.add_paragraph(str(ent), style='List Bullet')

        # Evidence Assessment
        assess_wrapper = (data or {}).get('assessment') or None
        doc.add_heading('Evidence Assessment', level=2)
        if assess_wrapper and assess_wrapper.get('assessment_successful'):
            assess = assess_wrapper.get('assessment') or {}
            overall = assess.get('overall_consistency_summary')
            if overall:
                doc.add_paragraph(f"Overall Summary: {overall}")
            cps = assess.get('consistent_points') or []
            if cps:
                doc.add_paragraph('Consistent Points:')
                for pt in cps:
                    doc.add_paragraph(pt, style='List Bullet')
            confs = assess.get('conflicting_points') or []
            if confs:
                doc.add_paragraph('Conflicting Points:')
                for pt in confs:
                    doc.add_paragraph(pt, style='List Bullet')
            sqs = assess.get('source_quality_assessments') or []
            if sqs:
                doc.add_paragraph('Source Quality Assessments:')
                for sq in sqs:
                    doc.add_paragraph(f"Source: {sq.get('source_id', 'N/A')}")
                    doc.add_paragraph(f"Relevance: {sq.get('relevance', 'N/A')}")
                    doc.add_paragraph(f"Trustworthiness: {sq.get('trustworthiness', 'N/A')}")
                    doc.add_paragraph(f"Timeliness: {sq.get('timeliness', 'N/A')}")
                    doc.add_paragraph(f"Authority: {sq.get('authority', 'N/A')}")
                    justification = sq.get('justification')
                    if justification:
                        doc.add_paragraph(f"Justification: {justification}")
        else:
            msg = 'Assessment unavailable or failed.'
            if isinstance(assess_wrapper, dict):
                msg = assess_wrapper.get('error') or msg
            doc.add_paragraph(msg)

        # Retrieved Context
        doc.add_heading('Retrieved Context', level=2)
        context_items = (data or {}).get('context') or []
        citation_index = ((data or {}).get('reasoned_answer') or {}).get('citation_index') or []
        srcfile_to_id: Dict[str, str] = {entry.get('source_file'): entry.get('source_id') for entry in citation_index if entry}

        for item in context_items:
            md = item.get('metadata') or {}
            file = md.get('source') or 'Unknown Source'
            source_id = item.get('source_id') or srcfile_to_id.get(file) or 'Source_?'
            page = md.get('page_number') or 'N/A'
            title = md.get('title') or 'No Title'
            content = item.get('text') or item.get('content') or ''
            doc.add_paragraph(f"[{source_id}, {title}, Page {page}] {content[:1000]}")

        # Style tweaks for readability
        for p in doc.paragraphs:
            for run in p.runs:
                run.font.size = Pt(11)

        return doc


query_export_service = QueryExportService()