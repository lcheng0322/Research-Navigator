import logging
from typing import List, Dict, Any, Optional, cast
import datetime
import json

from bertopic import BERTopic
from sentence_transformers import SentenceTransformer
from sklearn.feature_extraction.text import CountVectorizer
from celery.result import AsyncResult

from .llm_service import get_llm_response
from .vector_store_service import vector_store
from ..core.config import settings
from ..core.cache import redis_client
from ..worker import celery_app
from docx import Document
from docx.shared import Pt

# Custom stop words list for academic literature
ACADEMIC_STOP_WORDS = [
    'introduction', 'background', 'method', 'results', 'discussion', 'conclusion',
    'abstract', 'summary', 'references', 'study', 'paper', 'research', 'article',
    'data', 'model', 'system', 'et', 'al', 'however', 'therefore'
]

class ResearchGapAnalyzerService:
    """
    A service to analyze the knowledge base for research gaps, trends, and opportunities.
    This service now uses a Celery task for asynchronous execution.
    """
    CACHE_KEY_PREFIX = "gap_analysis_"
    CACHE_EXPIRATION = 3600 * 24  # 24 hours

    def __init__(self):
        logging.info("Initializing ResearchGapAnalyzerService...")
        try:
            from sklearn.feature_extraction.text import ENGLISH_STOP_WORDS
            stop_words = list(ENGLISH_STOP_WORDS) + ACADEMIC_STOP_WORDS
        except ImportError:
            stop_words = ACADEMIC_STOP_WORDS

        self.vectorizer_model = CountVectorizer(stop_words=stop_words, ngram_range=(1, 3))
        self.topic_model = BERTopic(
            embedding_model=self._load_embedding_model(),
            vectorizer_model=self.vectorizer_model
        )

    @staticmethod
    def _load_embedding_model():
        """Load the SentenceTransformer model for BERTopic, preferring local cache."""
        try:
            logging.info("Loading BERTopic embedding model from local cache...")
            return SentenceTransformer(settings.EMBEDDING_MODEL_NAME, local_files_only=True)
        except Exception:
            logging.info("Model not found locally, downloading from HuggingFace Hub...")
            return SentenceTransformer(settings.EMBEDDING_MODEL_NAME)

    def trigger_analysis(self, collection_names: Optional[List[str]] = None) -> str:
        """
        Triggers an asynchronous gap analysis task.
        Checks cache first before starting a new task.
        """
        if collection_names is None:
            collection_names = [settings.CHUNKS_COLLECTION_NAME, settings.SUMMARIES_COLLECTION_NAME]
        
        cache_key = f"{self.CACHE_KEY_PREFIX}{'_'.join(sorted(collection_names))}"
        cached_result = redis_client.get(cache_key)
        if cached_result:
            logging.info(f"Returning cached result for gap analysis: {cache_key}")
            return json.loads(cached_result)

        logging.info(f"No cache found. Triggering new gap analysis task for collections: {collection_names}")
        task = cast(AsyncResult, getattr(perform_gap_analysis, 'delay')(collection_names))
        return {"task_id": task.id, "status": "PENDING", "message": "Analysis has been started."}

    def get_analysis_result(self, task_id: str) -> Dict[str, Any]:
        """Retrieves the result of an analysis task."""
        task = celery_app.AsyncResult(task_id)
        if task.state == 'SUCCESS':
            return {"task_id": task_id, "status": task.state, "result": task.result}
        elif task.state == 'FAILURE':
            return {"task_id": task_id, "status": task.state, "error": str(task.info)}
        else:
            return {"task_id": task_id, "status": task.state}

    def export_gap_analysis_to_docx(self, analysis: Dict[str, Any]) -> Document:
        """
        Export the research gap analysis result to a DOCX document.
        The document includes summary statistics, core topics, outliers, trends info, and suggestions.
        """
        doc = Document()

        # Title
        doc.add_heading('Research Gap Analysis Report', level=0)

        # Summary section
        doc.add_heading('Summary', level=1)
        summary_lines = [
            f"Total Documents Analyzed: {analysis.get('total_documents_analyzed', 'N/A')}",
            f"Core Topics Found: {analysis.get('total_topics_found', 'N/A')}",
            f"Potential Gaps (Outliers): {analysis.get('outlier_documents_count', 0)}",
        ]
        for line in summary_lines:
            p = doc.add_paragraph(line)
            p.runs[0].font.size = Pt(11)

        # Research direction suggestion
        suggestion = analysis.get('research_gap_suggestion')
        if suggestion:
            doc.add_heading('Research Direction Suggestion', level=1)
            p = doc.add_paragraph(suggestion)
            p.runs[0].font.size = Pt(11)

        # Core topics table (skip topic == -1 which is noise)
        topics = analysis.get('topics', [])
        core_topics = [t for t in topics if t.get('Topic') != -1]
        if core_topics:
            doc.add_heading('Core Research Themes', level=1)
            table = doc.add_table(rows=1, cols=3)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Topic ID'
            hdr_cells[1].text = 'Topic Keywords'
            hdr_cells[2].text = 'Document Count'
            for t in core_topics[:20]:
                row_cells = table.add_row().cells
                row_cells[0].text = str(t.get('Topic', ''))
                # Format topic name similar to frontend (remove leading id if exists)
                name = str(t.get('Name', ''))
                formatted = ' '.join(word.capitalize() for word in name.split('_')[1:]) or name
                row_cells[1].text = formatted
                row_cells[2].text = str(t.get('Count', ''))

        # Outliers section
        outliers = analysis.get('outlier_documents', [])
        if outliers:
            doc.add_heading('Potential Research Gaps (Outliers)', level=1)
            doc.add_paragraph(
                'The following documents were identified as outliers because their content does not fit into any of the main research themes. They may represent niche topics, new concepts, or potential research gaps.'
            )
            for i, doc_item in enumerate(outliers[:10], start=1):
                meta = doc_item.get('metadata', {})
                source = meta.get('source', 'Unknown')
                page = meta.get('page_number')
                text = (doc_item.get('text') or '')
                snippet = text[:500] + ('...' if len(text) > 500 else '')
                entry = f"{i}. [{source}{', Page ' + str(page) if page is not None else ''}]\n{snippet}"
                doc.add_paragraph(entry)

        # Trends section (no charts; inform availability)
        trends = analysis.get('trends')
        doc.add_heading('Topic Trend Analysis', level=1)
        if trends and isinstance(trends, list) and len(trends) > 0:
            doc.add_paragraph(f"Trend data available with {len(trends)} records spanning time.")
            doc.add_paragraph("Refer to the application for interactive charts.")
        else:
            doc.add_paragraph('No trend data available.')

        return doc

@celery_app.task(name="perform_gap_analysis")
def perform_gap_analysis(collection_names: List[str]) -> Dict[str, Any]:
    """
    Celery task to perform the computationally intensive topic modeling and analysis.
    """
    service = ResearchGapAnalyzerService() # Re-instantiate service within the task
    cache_key = f"{service.CACHE_KEY_PREFIX}{'_'.join(sorted(collection_names))}"

    logging.info(f"[Task] Starting research gap analysis on collections: {collection_names}")
    
    # 1. Fetch and prepare data
    all_docs_with_meta = []
    for name in collection_names:
        docs = vector_store.get_all_documents(collection_name=name)
        all_docs_with_meta.extend(docs)

    if not all_docs_with_meta:
        raise ValueError("No documents found in any collection to perform analysis.")

    narrative_docs = [doc for doc in all_docs_with_meta if 'source' in doc.get('metadata', {}) and not doc['metadata']['source'].endswith('.csv')]
    if not narrative_docs:
        raise ValueError("No narrative documents available for analysis.")

    texts = [doc['text'] for doc in narrative_docs]
    timestamps = [doc['metadata'].get('publication_year') for doc in narrative_docs]
    
    # Filter out data without a valid year for trend analysis
    docs_for_trends = [(text, ts) for text, ts in zip(texts, timestamps) if isinstance(ts, int) and 1900 < ts <= datetime.date.today().year]
    trend_texts, trend_timestamps = zip(*docs_for_trends) if docs_for_trends else ([], [])

    # 2. Fit BERTopic model
    logging.info(f"[Task] Fitting BERTopic model on {len(texts)} documents...")
    topics, _ = service.topic_model.fit_transform(texts)
    topic_info = service.topic_model.get_topic_info()

    # 3. Perform Trend Analysis
    trends = None
    if trend_texts:
        try:
            logging.info(f"[Task] Performing trend analysis on {len(trend_texts)} documents...")
            topics_over_time = service.topic_model.topics_over_time(docs=list(trend_texts), timestamps=list(trend_timestamps))
            trends = topics_over_time.to_dict('records')
        except Exception as e:
            logging.warning(f"[Task] Could not generate topic trends: {e}")

    # 4. Generate Direction Suggestions for Outliers
    # Collect outlier documents (topic == -1) with text and minimal metadata for UI
    outlier_docs = [
        {
            "text": narrative_docs[i].get("text", ""),
            "metadata": {
                "source": narrative_docs[i].get("metadata", {}).get("source", "Unknown"),
                "page_number": narrative_docs[i].get("metadata", {}).get("page_number"),
            },
        }
        for i, topic in enumerate(topics)
        if topic == -1 and i < len(narrative_docs)
    ]
    direction_suggestion = "No significant outliers found to suggest new research directions."
    if outlier_docs:
        logging.info(f"[Task] Generating direction suggestions from {len(outlier_docs)} outlier documents...")
        # Summarize top outliers' text for LLM suggestion
        outlier_summary = "\n".join([f"- {doc.get('text', '')[:300]}..." for doc in outlier_docs[:10]])
        prompt = f'''
        You are a research strategist. The following text snippets are outliers from a topic modeling analysis of a large document set. This means they do not fit into any of the main identified themes.
        Analyze these outlier snippets to identify potential nascent trends or unexplored research gaps.

        OUTLIER SNIPPETS:
        {outlier_summary}

        Based on this, provide a concise summary of potential research directions or novel ideas suggested by these outliers.
        '''
        try:
            direction_suggestion = get_llm_response(prompt, use_reasoner=True)
        except Exception as e:
            logging.warning(f"[Task] Failed to get LLM suggestion for outliers: {e}")
            direction_suggestion = "Failed to generate suggestion due to an error."

    # 5. Compile final result
    result = {
        "total_documents_analyzed": len(set(doc['metadata']['source'] for doc in narrative_docs if 'source' in doc.get('metadata', {}))),
        "total_topics_found": len(topic_info[topic_info.Topic != -1]),
        "topics": topic_info.to_dict('records'),
        "trends": trends,
        "research_gap_suggestion": direction_suggestion,
        # Fields expected by frontend GapAnalysisPage
        "outlier_documents_count": len(outlier_docs),
        "outlier_documents": outlier_docs,
    }

    # Cache the final result
    redis_client.set(cache_key, json.dumps(result), ex=service.CACHE_EXPIRATION)
    logging.info(f"[Task] Gap analysis finished and result cached.")
    return result

# Singleton instance
research_gap_analyzer_service = ResearchGapAnalyzerService()
